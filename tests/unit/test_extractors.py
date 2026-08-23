from io import BytesIO

from knowledgeforge.ingestion.extract_docx import extract_docx
from knowledgeforge.ingestion.extract_markdown import extract_markdown


def test_extract_markdown_normalizes_text() -> None:
    pages = extract_markdown(BytesIO(b"# Heading\n\nA paragraph."))

    assert any("A paragraph." in text for _, text in pages)


def test_extract_docx_normalizes_paragraphs() -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    output = BytesIO()
    document.save(output)
    output.seek(0)

    assert extract_docx(output) == [(1, "First paragraph"), (2, "Second paragraph")]
